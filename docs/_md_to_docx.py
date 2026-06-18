"""
_md_to_docx.py — render the PCMMD instruction markdown files to styled .docx.
Handles: # headings, fenced ```code``` blocks (monospace + grey shade),
- bullets, | tables |, **bold**, `inline code`, blockquotes, hr.
Usage: python docs/_md_to_docx.py docs/INSTRUCTIONS_ABRAR.md docs/INSTRUCTIONS_TANJID.md
"""
import re
import sys

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.shared import Pt, RGBColor


def shade(paragraph, fill="F2F2F2"):
    pPr = paragraph._p.get_or_add_pPr()
    sh = OxmlElement("w:shd")
    sh.set(qn("w:val"), "clear"); sh.set(qn("w:fill"), fill)
    pPr.append(sh)


def add_runs(paragraph, text):
    """Render **bold** and `inline code` inside a paragraph."""
    for part in re.split(r"(\*\*.+?\*\*|`.+?`)", text):
        if not part:
            continue
        if part.startswith("**") and part.endswith("**"):
            r = paragraph.add_run(part[2:-2]); r.bold = True
        elif part.startswith("`") and part.endswith("`"):
            r = paragraph.add_run(part[1:-1]); r.font.name = "Consolas"; r.font.size = Pt(9.5)
            r.font.color.rgb = RGBColor(0xC7, 0x25, 0x4E)
        else:
            paragraph.add_run(part)


def code_block(doc, lines):
    for ln in lines:
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Pt(8)
        p.paragraph_format.space_after = Pt(0); p.paragraph_format.space_before = Pt(0)
        shade(p)
        r = p.add_run(ln if ln else " ")
        r.font.name = "Consolas"; r.font.size = Pt(9.5)


def convert(md_path, docx_path):
    lines = open(md_path, encoding="utf-8").read().split("\n")
    doc = Document()
    doc.styles["Normal"].font.name = "Calibri"
    doc.styles["Normal"].font.size = Pt(10.5)

    i, n = 0, len(lines)
    while i < n:
        line = lines[i]
        if line.strip().startswith("```"):           # fenced code
            i += 1; buf = []
            while i < n and not lines[i].strip().startswith("```"):
                buf.append(lines[i]); i += 1
            code_block(doc, buf); i += 1; continue
        if line.startswith("#"):                       # heading
            lvl = len(line) - len(line.lstrip("#"))
            doc.add_heading(line.lstrip("# ").strip(), level=min(lvl, 4)); i += 1; continue
        if line.strip().startswith("|") and "|" in line[1:]:   # table block
            tbl_lines = []
            while i < n and lines[i].strip().startswith("|"):
                tbl_lines.append(lines[i]); i += 1
            rows = [[c.strip() for c in r.strip().strip("|").split("|")] for r in tbl_lines
                    if not re.match(r"^\s*\|[\s:|-]+\|\s*$", r)]
            if rows:
                t = doc.add_table(rows=len(rows), cols=len(rows[0])); t.style = "Light Grid Accent 1"
                for ri, row in enumerate(rows):
                    for ci, cell in enumerate(row):
                        if ci < len(t.rows[ri].cells):
                            cp = t.rows[ri].cells[ci].paragraphs[0]; add_runs(cp, cell)
                            if ri == 0:
                                for rr in cp.runs:
                                    rr.bold = True
            continue
        if re.match(r"^\s*[-*]\s+", line):             # bullet
            p = doc.add_paragraph(style="List Bullet"); add_runs(p, re.sub(r"^\s*[-*]\s+", "", line)); i += 1; continue
        if line.strip().startswith(">"):               # blockquote
            p = doc.add_paragraph(); shade(p, "EAF2FB")
            add_runs(p, line.strip().lstrip(">").strip()); i += 1; continue
        if line.strip() in ("---", "***", "___"):       # hr
            doc.add_paragraph("_" * 60); i += 1; continue
        if line.strip() == "":
            doc.add_paragraph(); i += 1; continue
        p = doc.add_paragraph(); add_runs(p, line); i += 1

    doc.save(docx_path)
    print("wrote", docx_path)


if __name__ == "__main__":
    for md in sys.argv[1:]:
        convert(md, md.rsplit(".", 1)[0] + ".docx")
